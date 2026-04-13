"""Build docs/RAISE-User-Manual.docx from docs/RAISE-User-Manual.md (requires python-docx)."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt


def is_table_separator(line: str) -> bool:
    s = line.strip()
    if not s.startswith("|"):
        return False
    inner = s.strip("|").replace(" ", "")
    return bool(re.match(r"^:?-+:?(?:\|:?-+:?)*$", inner))


def parse_table_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    md_path = root / "docs" / "RAISE-User-Manual.md"
    out_path = root / "docs" / "RAISE-User-Manual.docx"

    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i]
        raw = line.rstrip("\n")
        s = raw.strip()

        if not s:
            doc.add_paragraph("")
            i += 1
            continue

        if s.startswith("# ") and not s.startswith("## "):
            doc.add_heading(s[2:].strip(), 0)
            i += 1
            continue
        if s.startswith("## ") and not s.startswith("### "):
            doc.add_heading(s[3:].strip(), 1)
            i += 1
            continue
        if s.startswith("### "):
            doc.add_heading(s[4:].strip(), 2)
            i += 1
            continue

        if s.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            rows: list[list[str]] = [parse_table_row(lines[i].strip())]
            i += 2  # header + separator
            while i < len(lines) and lines[i].strip().startswith("|"):
                row_line = lines[i].strip()
                if is_table_separator(row_line):
                    i += 1
                    continue
                rows.append(parse_table_row(row_line))
                i += 1
            if rows:
                cols = max(len(r) for r in rows)
                t = doc.add_table(rows=len(rows), cols=cols)
                t.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci in range(cols):
                        cell_text = row[ci] if ci < len(row) else ""
                        t.rows[ri].cells[ci].text = cell_text
            continue

        if s.startswith("|---") or is_table_separator(s):
            i += 1
            continue

        doc.add_paragraph(raw)
        i += 1

    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
